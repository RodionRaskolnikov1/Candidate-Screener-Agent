import json
import os
import sqlite3
from datetime import datetime
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
BASE_DIR        = Path(__file__).parent
RESUMES_DIR     = BASE_DIR / "resumes"
REPORTS_DIR     = BASE_DIR / "reports"
DB_PATH         = BASE_DIR / "screening_results.db"
JD_PATH         = BASE_DIR / "job_description.txt"
MODEL           = "meta-llama/llama-4-scout-17b-16e-instruct"

RESUMES_DIR.mkdir(exist_ok=True)
REPORTS_DIR.mkdir(exist_ok=True)

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

# ---------------------------------------------------------------------------
# DATABASE SETUP
# ---------------------------------------------------------------------------

def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS candidates (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            filename        TEXT NOT NULL,
            candidate_name  TEXT,
            fit_score       INTEGER,
            recommendation  TEXT,
            matched_skills  TEXT,
            missing_skills  TEXT,
            summary         TEXT,
            screened_at     TEXT
        )
    """)
    conn.commit()
    conn.close()

# ---------------------------------------------------------------------------
# TOOL FUNCTIONS
# ---------------------------------------------------------------------------

def extract_resume_text(pdf_path: str) -> str:
    """
    Extracts all text from a resume PDF file.
    Returns the raw text content, or an error message if extraction fails.
    """
    try:
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text.strip())
        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            return "ERROR: Could not extract any text from this PDF. It may be image-based."
        return full_text[:6000]  # cap at 6000 chars to stay within token limits
    except Exception as e:
        return f"ERROR: Failed to read PDF: {e}"


def evaluate_candidate(resume_text: str, job_description: str, filename: str) -> dict:
    """
    Sends resume text + job description to the LLM.
    Returns a structured evaluation as a Python dict with keys:
      candidate_name, fit_score (0-10), recommendation,
      matched_skills (list), missing_skills (list), summary
    """
    prompt = f"""
You are an expert HR screening assistant. Evaluate the following resume against the job description.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Return ONLY a valid JSON object with these exact keys:
{{
  "candidate_name": "Full name from resume, or 'Unknown' if not found",
  "fit_score": <integer 0-10>,
  "recommendation": "Shortlist" or "Reject" or "Maybe",
  "matched_skills": ["skill1", "skill2", ...],
  "missing_skills": ["skill1", "skill2", ...],
  "summary": "2-3 sentence honest summary of the candidate's fit"
}}

Scoring guide:
  8-10 = Strong fit, meets most requirements
  5-7  = Partial fit, meets some requirements
  0-4  = Poor fit, missing critical requirements

Recommendation guide:
  Shortlist = score >= 7
  Maybe     = score 5-6
  Reject    = score <= 4

Return only the JSON, no explanation, no markdown fences.
"""
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()
        # strip markdown fences if model added them anyway
        raw = raw.replace("```json", "").replace("```", "").strip()
        result = json.loads(raw)
        result["filename"] = filename
        return result
    except json.JSONDecodeError as e:
        return {
            "filename": filename,
            "candidate_name": "Parse Error",
            "fit_score": 0,
            "recommendation": "Reject",
            "matched_skills": [],
            "missing_skills": [],
            "summary": f"Could not parse LLM response: {e}",
        }
    except Exception as e:
        return {
            "filename": filename,
            "candidate_name": "Error",
            "fit_score": 0,
            "recommendation": "Reject",
            "matched_skills": [],
            "missing_skills": [],
            "summary": f"Evaluation failed: {e}",
        }


def save_to_db(result: dict) -> str:
    """
    Saves a candidate evaluation result to the SQLite database.
    Returns confirmation string.
    """
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("""
            INSERT INTO candidates
              (filename, candidate_name, fit_score, recommendation,
               matched_skills, missing_skills, summary, screened_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            result.get("filename", "unknown"),
            result.get("candidate_name", "Unknown"),
            result.get("fit_score", 0),
            result.get("recommendation", "Reject"),
            json.dumps(result.get("matched_skills", [])),
            json.dumps(result.get("missing_skills", [])),
            result.get("summary", ""),
            datetime.now().isoformat(),
        ))
        conn.commit()
        conn.close()
        return f"Saved: {result.get('candidate_name')} (score: {result.get('fit_score')}/10)"
    except Exception as e:
        return f"DB error: {e}"


def generate_report(job_title: str = "Open Position") -> str:
    """
    Reads all candidates from DB, ranks them by fit score,
    and generates a formatted .docx report in the reports/ folder.
    Returns the path to the generated report.
    """
    try:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute("""
            SELECT candidate_name, fit_score, recommendation,
                   matched_skills, missing_skills, summary, filename, screened_at
            FROM candidates
            ORDER BY fit_score DESC
        """).fetchall()
        conn.close()

        if not rows:
            return "No candidates in DB. Run screening first."

        doc = Document()

        # --- Page setup ---
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin   = Inches(1)
        section.top_margin  = section.bottom_margin  = Inches(1)

        # --- Title ---
        title = doc.add_heading(f"Candidate Screening Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size    = Pt(20)
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        subtitle = doc.add_paragraph(f"Position: {job_title}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size  = Pt(12)
        subtitle.runs[0].font.bold  = True

        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size  = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # --- Summary stats ---
        total       = len(rows)
        shortlisted = sum(1 for r in rows if r[2] == "Shortlist")
        maybe       = sum(1 for r in rows if r[2] == "Maybe")
        rejected    = sum(1 for r in rows if r[2] == "Reject")

        stats = doc.add_heading("Summary", level=2)
        stats.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        hdr = summary_table.rows[0].cells
        for cell, label, value in zip(
            hdr,
            ["Total Screened", "Shortlisted", "Maybe", "Rejected"],
            [total, shortlisted, maybe, rejected],
        ):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}\n{value}")
            run.font.bold = True
            run.font.size = Pt(11)

        doc.add_paragraph("")

        # --- Per candidate sections ---
        doc.add_heading("Candidate Evaluations", level=2).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        for i, row in enumerate(rows, 1):
            name, score, recommendation, matched_raw, missing_raw, summary, filename, screened_at = row
            matched = json.loads(matched_raw) if matched_raw else []
            missing = json.loads(missing_raw) if missing_raw else []

            # Recommendation color
            rec_colors = {
                "Shortlist": RGBColor(0x1E, 0x8B, 0x4C),
                "Maybe":     RGBColor(0xD4, 0x7F, 0x00),
                "Reject":    RGBColor(0xC0, 0x39, 0x2B),
            }
            rec_color = rec_colors.get(recommendation, RGBColor(0x33, 0x33, 0x33))

            # Candidate heading
            heading = doc.add_heading(f"{i}. {name}", level=3)
            heading.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Score + recommendation line
            score_para = doc.add_paragraph()
            score_para.add_run("Fit Score: ").font.bold = True
            score_run = score_para.add_run(f"{score}/10")
            score_run.font.bold = True
            score_run.font.size = Pt(13)
            score_para.add_run("    ")
            score_para.add_run("Recommendation: ").font.bold = True
            rec_run = score_para.add_run(recommendation)
            rec_run.font.bold  = True
            rec_run.font.color.rgb = rec_color

            # Summary
            sum_para = doc.add_paragraph()
            sum_para.add_run("Summary: ").font.bold = True
            sum_para.add_run(summary)

            # Matched skills
            if matched:
                matched_para = doc.add_paragraph()
                matched_para.add_run("Matched Skills: ").font.bold = True
                matched_para.add_run(", ".join(matched))

            # Missing skills
            if missing:
                missing_para = doc.add_paragraph()
                mr = missing_para.add_run("Missing Skills: ")
                mr.font.bold = True
                missing_para.add_run(", ".join(missing))

            # File reference
            file_para = doc.add_paragraph()
            fr = file_para.add_run(f"Resume file: {filename}")
            fr.font.size  = Pt(9)
            fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Divider (empty paragraph with bottom border)
            div = doc.add_paragraph()
            from docx.oxml.ns import qn
            from docx.oxml    import OxmlElement
            pPr = div._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # --- Save ---
        timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = REPORTS_DIR / f"screening_report_{timestamp}.docx"
        doc.save(report_path)
        return f"Report saved: {report_path}"

    except Exception as e:
        return f"Report generation failed: {e}"


# ---------------------------------------------------------------------------
# TOOL DEFINITIONS (for the LLM)
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "extract_resume_text",
            "description": "Extracts text from a resume PDF file. Pass the full file path.",
            "parameters": {
                "type": "object",
                "properties": {
                    "pdf_path": {"type": "string", "description": "Full path to the PDF file."}
                },
                "required": ["pdf_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "evaluate_candidate",
            "description": "Evaluates a resume against a job description. Returns structured JSON with score and recommendation.",
            "parameters": {
                "type": "object",
                "properties": {
                    "resume_text":       {"type": "string", "description": "Raw text extracted from the resume."},
                    "job_description":   {"type": "string", "description": "Full job description text."},
                    "filename":          {"type": "string", "description": "Original PDF filename for reference."},
                },
                "required": ["resume_text", "job_description", "filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_to_db",
            "description": "Saves a candidate evaluation result dict to the SQLite database.",
            "parameters": {
                "type": "object",
                "properties": {
                    "result": {"type": "object", "description": "The evaluation result dict from evaluate_candidate."}
                },
                "required": ["result"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_report",
            "description": "Reads all candidates from DB, ranks them, and generates a .docx report.",
            "parameters": {
                "type": "object",
                "properties": {
                    "job_title": {"type": "string", "description": "Job title to display in the report header."}
                },
                "required": [],
            },
        },
    },
]

TOOL_MAP = {
    "extract_resume_text": extract_resume_text,
    "evaluate_candidate":  evaluate_candidate,
    "save_to_db":          save_to_db,
    "generate_report":     generate_report,
}


# ---------------------------------------------------------------------------
# AGENT LOOP (pure SDK — same pattern as Day 14)
# ---------------------------------------------------------------------------

def execute_tool(name: str, args: dict):
    func = TOOL_MAP.get(name)
    if not func:
        return f"Unknown tool: {name}"
    try:
        return func(**args)
    except Exception as e:
        return f"Tool error: {e}"


def run_screening():
    # Load job description
    if not JD_PATH.exists():
        print(f"ERROR: {JD_PATH} not found. Create this file with the job description.")
        return

    job_description = JD_PATH.read_text(encoding="utf-8").strip()
    if not job_description:
        print("ERROR: job_description.txt is empty.")
        return

    # Find resume PDFs
    pdf_files = list(RESUMES_DIR.glob("*.pdf"))
    if not pdf_files:
        print(f"ERROR: No PDF files found in {RESUMES_DIR}")
        return

    # Extract job title (first non-empty line of JD)
    job_title = next(
        (line.strip() for line in job_description.splitlines() if line.strip()),
        "Open Position"
    )

    print(f"\n{'='*60}")
    print(f"  Candidate Screener Agent")
    print(f"  Position : {job_title}")
    print(f"  Resumes  : {len(pdf_files)} found")
    print(f"{'='*60}\n")

    # Build the instruction for the agent
    pdf_paths_str = "\n".join(str(p) for p in pdf_files)
    user_message = f"""
Screen the following resume PDFs against the provided job description.

For EACH resume:
1. Call extract_resume_text with the PDF path
2. Call evaluate_candidate with the extracted text, job description, and filename
3. Call save_to_db with the evaluation result

After ALL resumes are processed, call generate_report with job_title="{job_title}".

Resume files to process:
{pdf_paths_str}

Job description:
{job_description}
"""

    conversation_history = [{"role": "user", "content": user_message}]

    system_prompt = (
        "You are a candidate screening assistant. "
        "Process each resume file one by one using the tools provided. "
        "Always extract text first, then evaluate, then save. "
        "After all resumes are done, generate the report. "
        "Be thorough and process every single file listed."
    )

    print("Starting screening...\n")

    # Agent loop
    while True:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "system", "content": system_prompt}] + conversation_history,
            tools=TOOL_DEFINITIONS,
            tool_choice="auto",
        )

        message = response.choices[0].message

        if message.tool_calls:
            conversation_history.append({
                "role": "assistant",
                "content": message.content,
                "tool_calls": [
                    {
                        "id": tc.id,
                        "type": "function",
                        "function": {
                            "name": tc.function.name,
                            "arguments": tc.function.arguments,
                        },
                    }
                    for tc in message.tool_calls
                ],
            })

            for tc in message.tool_calls:
                tool_name = tc.function.name
                tool_args = json.loads(tc.function.arguments)

                print(f"  → {tool_name}({list(tool_args.keys())})")
                result = execute_tool(tool_name, tool_args)
                print(f"    {result if isinstance(result, str) else 'OK'}\n")

                conversation_history.append({
                    "role": "tool",
                    "tool_call_id": tc.id,
                    "content": json.dumps(result) if not isinstance(result, str) else result,
                })

            continue

        # Final message — agent is done
        print(f"\n{'='*60}")
        print("Agent complete.")
        print(f"{'='*60}")
        print(f"\n{message.content}")
        break


# ---------------------------------------------------------------------------
# ENTRY POINT
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    init_db()
    run_screening()